import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from io import StringIO, BytesIO
from docx import Document
from docx.shared import Inches

st.title("БЖБ және ТЖБ жұмыстарының нәтижелерін талдау")
st.write("19 жалпы білім беретін мектеп КММ")

csv_text = st.text_area("Excel-ден алынған CSV мәтін түріндегі кестені енгізіңіз", height=200)

if csv_text.strip():
    try:
        # ----------------------------------------------------
        # CSV ЖҮКТЕУ
        # ----------------------------------------------------
        df = pd.read_csv(StringIO(csv_text))

        # Проценттерді санға ауыстыру
        for col in df.columns:
            if df[col].astype(str).str.contains("%").any():
                df[col] = (
                    df[col]
                    .astype(str)
                    .str.replace("%", "")
                    .str.replace(",", ".")
                    .str.strip()
                    .astype(float)
                )

        st.success("Кесте жүктелді!")
        st.dataframe(df)

        # ----------------------------------------------------
        # ҚАЖЕТТІ БАҒАНДАРДЫ ТАБУ
        # ----------------------------------------------------
        quality_col = None
        success_col = None

        for col in df.columns:
            col_low = col.lower()
            if "сапа" in col_low:
                quality_col = col
            if "үлгер" in col_low:
                success_col = col

        if not quality_col or not success_col:
            st.error("«Білім Сапасы (%)» немесе «Үлгерім (%)» бағандары табылмады.")
            st.stop()

        st.info(f"Білім сапасы бағаны: **{quality_col}**")
        st.info(f"Үлгерім бағаны: **{success_col}**")

        # ----------------------------------------------------
        # ҰСЫНЫСТАР (Қайталайтын тақырыптар)
        # ----------------------------------------------------
        st.subheader("📌 Ұсыныстар (Қайталау қажет тақырыптар)")

        recommendations = {
            "БЖБ 1": "Кинематика: Қозғалыс теңдеулері, графиктерді талдау (жылдамдық, үдеу).",
            "БЖБ 2": "Динамика және Импульс: Ньютон заңдары, күштер, импульстің сақталу заңы.",
            "ТЖБ": "Кешенді есептер: Кинематика мен динамиканы біріктіретін есептер."
        }

        for assess, text in recommendations.items():
            st.markdown(f"**{assess}:** {text}")

        # ----------------------------------------------------
        # БАҒАЛАУ ТҮРЛЕРІ
        # ----------------------------------------------------
        assess_types = ["БЖБ 1", "БЖБ 2", "ТЖБ"]

        # Word құжат
        document = Document()
        document.add_heading("Бақылау жұмыстарының нәтижелері", level=1)
        image_buffers = []

        # ----------------------------------------------------
        # ДИАГРАММАЛАР ЖАСАУ
        # ----------------------------------------------------
        for assess in assess_types:
            subset = df[df["Бағалау түрі"].str.contains(assess, case=False, na=False)]

            if subset.empty:
                continue

            st.subheader(f"{assess}: Білім сапасы мен үлгерім көрсеткіші")

            labels = subset["Сынып"]
            q = subset[quality_col]
            u = subset[success_col]

            fig, ax = plt.subplots(figsize=(8, 4))
            x = range(len(labels))

            ax.bar([p - 0.2 for p in x], q, width=0.4, label="Білім сапасы")
            ax.bar([p + 0.2 for p in x], u, width=0.4, label="Үлгерім")

            # --- ДИАГРАММА ҮСТІНДЕ ПАЙЫЗДАР ---
            for i, val in enumerate(q):
                ax.text(i - 0.2, val + 1, f"{val}%", ha='center', va='bottom', fontsize=9)

            for i, val in enumerate(u):
                ax.text(i + 0.2, val + 1, f"{val}%", ha='center', va='bottom', fontsize=9)

            ax.set_xticks(x)
            ax.set_xticklabels(labels)
            ax.set_title(f"{assess}: Білім сапасы мен үлгерім")
            ax.set_ylabel("%")
            ax.legend()

            st.pyplot(fig)

            # Word үшін суреттерді сақтау
            img_buf = BytesIO()
            fig.savefig(img_buf, format="png", dpi=200)
            img_buf.seek(0)
            image_buffers.append((assess, img_buf))

        # ----------------------------------------------------
        # WORD ФАЙЛ ЖАЗУ
        # ----------------------------------------------------
        for title, img_buf in image_buffers:
            document.add_heading(title, level=2)
            document.add_picture(img_buf, width=Inches(6))

        output = BytesIO()
        document.save(output)
        output.seek(0)

        st.download_button(
            "📥 Word файлды жүктеу",
            data=output,
            file_name="Бақылау_жұмыстары_талдау.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    except Exception as e:
        st.error(f"Қате анықталды: {e}")
